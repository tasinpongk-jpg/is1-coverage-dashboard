"""Layer 2 of the Hermes MDA-FS harvester — download + extract + write to vault.

Reads data/harvest-queue.json (produced by harvest_filings.py), downloads each
pending MDA / FS filing from SET, extracts text, and delegates the markdown
write to scripts/vault_raw_writer.py (Loop 4 v5). State persists in
data/harvest-state.json to avoid re-downloading.

Why this is small: scripts/vault_raw_writer.py already does extraction
classification, frontmatter rendering, atomic writes, and SHA-256 dedup.
Layer 2's job is only to:
  1. Resolve the downloadUrl via /api/set/news/{id}/detail
  2. Fetch the ZIP / PDF bytes
  3. Hand them to vault_raw_writer.project_one() in the expected raw_markdown shape
  4. Persist the result back to harvest-queue.json

Usage:
    python scripts/harvest_download.py                    # process all pending
    python scripts/harvest_download.py --ticker TPAC      # one ticker only
    python scripts/harvest_download.py --limit 5          # first 5 items
    python scripts/harvest_download.py --dry-run          # don't write vault
"""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import io
import json
import os
import sys
import time
import zipfile
from pathlib import Path
from typing import Any

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO))
sys.path.insert(0, str(REPO / "surveillance"))
sys.path.insert(0, str(REPO / "scripts"))

from client import SetNewsClient, HEADERS  # noqa: E402

import vault_raw_writer  # noqa: E402
import enrich_filing as ef  # noqa: E402

DATA_DIR = REPO / "data"
QUEUE_FILE = DATA_DIR / "harvest-queue.json"
STATE_FILE = DATA_DIR / "harvest-state.json"

# Load ~/.hermes/.env-harvest if it exists (best-effort, silent on errors).
# This lets users drop their MINIMAX_API_KEY in a single place without
# needing to export it system-wide. The Windows task wrapper registered by
# scripts/register_harvest.ps1 ALSO loads this file, so the env is available
# even when launched by Task Scheduler (which doesn't inherit interactive
# shell env vars).
_HERMES_DOTENV = Path.home() / ".hermes" / ".env-harvest"
if _HERMES_DOTENV.exists():
    try:
        for _line in _HERMES_DOTENV.read_text(encoding="utf-8").splitlines():
            _line = _line.strip()
            if not _line or _line.startswith("#") or "=" not in _line:
                continue
            _k, _, _v = _line.partition("=")
            _k, _v = _k.strip(), _v.strip()
            # Don't overwrite values already set in the env.
            if _k and _k not in os.environ:
                os.environ[_k] = _v
    except (OSError, UnicodeDecodeError):
        pass

# Reuse the doc-extraction helpers from enrich_filing.py — they already
# handle DOCX/XLSX/PDF and nested ZIPs (from Loop 4 v5).
_DOC_FACTORIES = {
    "MDA":     ef._extract_text_from_pdf if hasattr(ef, "_extract_text_from_pdf") else None,
}


# ---------------------------------------------------------------- queue I/O

def load_queue() -> dict[str, Any]:
    if not QUEUE_FILE.exists():
        return {"generated": "", "items": {}}
    try:
        return json.loads(QUEUE_FILE.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {"generated": "", "items": {}}


def load_state() -> dict[str, Any]:
    if not STATE_FILE.exists():
        return {"completed": {}}
    try:
        return json.loads(STATE_FILE.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {"completed": {}}


def save_queue(queue: dict[str, Any]) -> None:
    queue["generated"] = dt.datetime.now(dt.timezone.utc).isoformat(timespec="seconds")
    QUEUE_FILE.write_text(
        json.dumps(queue, ensure_ascii=False, indent=2, sort_keys=False) + "\n",
        encoding="utf-8",
    )


def save_state(state: dict[str, Any]) -> None:
    STATE_FILE.write_text(
        json.dumps(state, ensure_ascii=False, indent=2, sort_keys=False) + "\n",
        encoding="utf-8",
    )


# ---------------------------------------------------------------- SET detail

def fetch_detail(client: SetNewsClient, news_id: str) -> dict[str, Any]:
    """GET /api/set/news/{news_id}/detail. Returns the JSON body."""
    # SetNewsClient doesn't expose a detail() method, so we use it as a context
    # and call the underlying client with the same throttle discipline.
    client.warmup()
    client._throttle()  # noqa: SLF001 — same discipline as the public API
    url = f"https://www.set.or.th/api/set/news/{news_id}/detail"
    r = client._client.get(url)  # noqa: SLF001
    r.raise_for_status()
    return r.json()


def download_bytes(client: SetNewsClient, url: str, timeout: float = 60.0) -> bytes:
    """Download a file with the same UA/headers; return raw bytes."""
    client.warmup()
    client._throttle()  # noqa: SLF001
    r = client._client.get(url, timeout=timeout)
    r.raise_for_status()
    return r.content


# ---------------------------------------------------------------- extraction

def extract_mda_pdf(pdf_bytes: bytes) -> tuple[str, int]:
    """Extract text from a single-PDF MDA. Returns (text, page_count).

    Detects scanned/encrypted PDFs (zero text extracted) and falls back to
    OCR via tesseract if available. If OCR is not installed, marks the
    filing as needs_review so the user can process it manually.
    """
    try:
        import pypdf  # type: ignore
    except ImportError:
        return "", 0
    try:
        reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
        pages: list[str] = []
        for p in reader.pages:
            try:
                pages.append(p.extract_text() or "")
            except Exception:  # noqa: BLE001
                pages.append("")
        text = "\n\n".join(pages).strip()
        if not text:
            # Scanned PDF or only-image PDF. Try OCR fallback.
            print(f"  [harvest-dl] PDF has 0 extractable text "
                  f"({len(reader.pages)} pages) — trying OCR",
                  flush=True)
            ocr_text = _ocr_pdf_fallback(pdf_bytes, len(reader.pages))
            if ocr_text:
                return ocr_text, len(reader.pages)
            print(f"  [harvest-dl] OCR unavailable or empty — "
                  f"mark as needs_review (set TESSERACT_CMD env var to enable)",
                  flush=True)
        return text, len(reader.pages)
    except Exception as exc:  # noqa: BLE001
        print(f"  [harvest-dl] MDA PDF parse error: {type(exc).__name__}: {exc}", flush=True)
        return "", 0


def _ocr_pdf_fallback(pdf_bytes: bytes, page_count: int) -> str:
    """OCR a scanned PDF using m3 vision (preferred) or tesseract (fallback).

    Strategy (priority order):
      1. m3 vision API — sends the PDF as base64, m3 does OCR + structure
         recognition in one shot. Requires MINIMAX_API_KEY env var.
         Cost: ~$0.01-0.05 per page. Most reliable for Thai + English mixed text.
      2. tesseract OCR — render via pypdfium2, shell out to tesseract binary.
         Requires tesseract installed at TESSERACT_CMD path or in PATH.
         Free, but quality is mediocre for Thai (eng model only).
      3. Return "" — caller marks as needs_review.

    Returns the extracted text (or "" if all paths fail).
    """
    # 1. m3 vision (preferred)
    text = _m3_vision_ocr(pdf_bytes)
    if text and text.strip():
        return text.strip()

    # 2. tesseract fallback
    text = _tesseract_ocr(pdf_bytes)
    if text and text.strip():
        return text.strip()

    return ""


def _m3_vision_ocr(pdf_bytes: bytes) -> str:
    """Use m3 vision (Anthropic Messages API) to OCR a PDF.

    Mirrors the request shape in scripts/enrich_filing.py. Requires
    MINIMAX_API_KEY env var. Returns "" if no key or call fails.
    """
    import base64
    import urllib.error
    import urllib.request

    api_key = os.environ.get("MINIMAX_API_KEY") or os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        return ""
    try:
        import enrich_filing as ef  # reuse constants
        m3_url = ef.M3_BASE_URL + "/v1/messages"
        m3_model = ef.M3_MODEL
        max_tokens = ef.M3_MAX_TOKENS
    except ImportError:
        return ""

    b64 = base64.standard_b64encode(pdf_bytes).decode("ascii")
    body = {
        "model": m3_model,
        "max_tokens": max_tokens,
        "messages": [{
            "role": "user",
            "content": [
                {"type": "document", "source": {
                    "type": "base64",
                    "media_type": "application/pdf",
                    "data": b64,
                }},
                {"type": "text", "text":
                 "Extract all readable text from this scanned PDF document. "
                 "Return only the extracted text, preserving section structure "
                 "with markdown headings where possible. If the document is in "
                 "Thai, return Thai text; if English, return English. Do not "
                 "summarize — extract verbatim."},
            ],
        }],
    }
    try:
        req = urllib.request.Request(
            m3_url,
            data=json.dumps(body).encode("utf-8"),
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "Content-Type": "application/json",
            },
        )
        with urllib.request.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode("utf-8"))
        # Extract text blocks from the response.
        parts = []
        for block in result.get("content", []):
            if block.get("type") == "text":
                parts.append(block.get("text", ""))
        return "\n\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        print(f"  [harvest-dl] m3 vision OCR error: {type(exc).__name__}: {exc}",
              flush=True)
        return ""


def _tesseract_ocr(pdf_bytes: bytes) -> str:
    """OCR a scanned PDF using tesseract (fallback if m3 unavailable).

    Requires tesseract binary installed. Renders each page via pypdfium2
    and writes a BMP file (no PIL dependency — works even when the venv's
    Pillow is broken). Tesseract accepts BMP input natively.

    Set TESSERACT_CMD env var to override the binary path; otherwise
    the binary must be on PATH.
    """
    import shutil
    import struct
    import subprocess
    import tempfile

    tesseract_cmd = os.environ.get("TESSERACT_CMD") or shutil.which("tesseract")
    if not tesseract_cmd:
        return ""
    try:
        import pypdfium2 as pdfium  # type: ignore
    except ImportError:
        return ""

    text_chunks: list[str] = []
    try:
        pdf = pdfium.PdfDocument(pdf_bytes)
        for page_idx in range(len(pdf)):
            tmp_bmp = None
            try:
                page = pdf[page_idx]
                bm = page.render(scale=2)
                width, height = bm.width, bm.height
                stride = bm.stride
                # Write a 24-bit BMP from the bitmap buffer. pypdfium2 may
                # return BGR (3 bytes/pixel) or BGRA (4 bytes/pixel) — we
                # detect via the stride and read the right number of bytes.
                # We avoid PIL entirely because the user's venv may have a
                # broken Pillow install (cp311 .pyd in cp314 venv). BMP is
                # trivially encoded and tesseract reads it natively.
                row_size = (width * 3 + 3) & ~3
                pixel_data_size = row_size * height
                file_size = 54 + pixel_data_size
                bmp = bytearray()
                bmp += b'BM'
                bmp += struct.pack('<I', file_size)
                bmp += struct.pack('<HH', 0, 0)
                bmp += struct.pack('<I', 54)
                bmp += struct.pack('<I', 40)
                bmp += struct.pack('<i', width)
                bmp += struct.pack('<i', height)
                bmp += struct.pack('<HH', 1, 24)
                bmp += struct.pack('<I', 0)
                bmp += struct.pack('<I', pixel_data_size)
                bmp += struct.pack('<ii', 2835, 2835)
                bmp += struct.pack('<II', 0, 0)
                # Convert BGR/BGRA -> BGR rows, BMP is bottom-up.
                buf = bytes(bm.buffer)
                bytes_per_pixel = stride // width if width > 0 else 3
                for y in range(height - 1, -1, -1):
                    row_start = y * stride
                    for x in range(width):
                        px = row_start + x * bytes_per_pixel
                        b = buf[px + 0]
                        g = buf[px + 1]
                        r = buf[px + 2]
                        bmp += bytes([b, g, r])
                    pad = row_size - width * 3
                    bmp += b'\x00' * pad
                with tempfile.NamedTemporaryFile(suffix=".bmp", delete=False) as tmp:
                    tmp.write(bytes(bmp))
                    tmp_bmp = tmp.name
                # Tesseract supports both eng and tha languages.
                result = subprocess.run(
                    [tesseract_cmd, tmp_bmp, "-", "-l", "eng+tha"],
                    capture_output=True, text=True, timeout=60,
                )
                if result.returncode == 0 and result.stdout.strip():
                    text_chunks.append(result.stdout)
            finally:
                if tmp_bmp:
                    try:
                        os.unlink(tmp_bmp)
                    except OSError:
                        pass
    except Exception as exc:  # noqa: BLE001
        print(f"  [harvest-dl] tesseract OCR error: {type(exc).__name__}: {exc}",
              flush=True)
        return ""
    return "\n\n".join(text_chunks).strip()


def extract_fs_zip(zip_bytes: bytes) -> list[dict[str, Any]]:
    """Decompose an FS ZIP into per-document text. Mirrors the Loop 4 v5 logic
    in scripts/fetch_set_financial_filings.py:extract_texts + pick_member.

    Returns list of {"doctype": "FS"|"NOTES"|"AUDITOR"|"OTHER",
                     "member_filename": str, "text": str, "raw_bytes_len": int,
                     "page_count"|"sheet_count": int}.
    """
    out: list[dict[str, Any]] = []
    members = list(_iter_zip_members(zip_bytes))
    for name, payload in members:
        suffix = Path(name).suffix.lower()
        text = ""
        status = "ok"
        page_count = None
        sheet_count = None
        doctype = "OTHER"  # default; overridden by successful extraction
        try:
            if suffix == ".pdf":
                text, page_count = extract_mda_pdf(payload)
                # Embedded PDFs inside a ZIP are usually NOTES (full statements),
                # not MDA — the SET convention is MDA ships as a standalone PDF.
                doctype = "NOTES"
            elif suffix == ".doc":
                text = _doc_text(payload)
                doctype = _classify_office_member(name, text)
            elif suffix == ".docx":
                text = _docx_text(payload)
                doctype = _classify_office_member(name, text)
            elif suffix == ".xls":
                text = _xls_text(payload)
                sheet_count = text.count("# Sheet:")
                doctype = _classify_office_member(name, text)
            elif suffix == ".xlsx":
                text = _xlsx_text(payload)
                sheet_count = text.count("# Sheet:")
                doctype = _classify_office_member(name, text)
            elif suffix in {".txt", ".csv", ".xml"}:
                text = payload.decode("utf-8", errors="replace")
                doctype = _classify_office_member(name, text)
            else:
                continue
        except Exception as exc:  # noqa: BLE001
            print(f"  [harvest-dl] FS member {name} parse error: {exc}", flush=True)
            status = "parse_error"
            # doctype stays "OTHER" so the doc is still recorded and can be skipped later
        out.append({
            "doctype": doctype,
            "member_filename": name,
            "text": text,
            "raw_bytes_len": len(payload),
            "extraction_status": status,
            "page_count": page_count,
            "sheet_count": sheet_count,
        })
    return out


def _iter_zip_members(data: bytes, prefix: str = "") -> list[tuple[str, bytes]]:
    out: list[tuple[str, bytes]] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                name = prefix + info.filename
                payload = zf.read(info)
                suffix = Path(name).suffix.lower()
                if suffix in {".doc", ".docx", ".xls", ".xlsx", ".pdf", ".txt", ".csv", ".xml"}:
                    out.append((name, payload))
                elif zipfile.is_zipfile(io.BytesIO(payload)):
                    out.extend(_iter_zip_members(payload, prefix=f"{name}/"))
    except zipfile.BadZipFile as exc:
        print(f"  [harvest-dl] BadZipFile: {exc}", flush=True)
    return out


def _doc_text(data: bytes) -> str:
    """Extract text from legacy .doc (binary Word format) via antiword CLI.

    Returns "" if antiword is not available or fails. Falls back silently.
    """
    import subprocess
    import tempfile
    try:
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(data)
            tmp_path = tmp.name
        result = subprocess.run(
            ["antiword", tmp_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        try:
            os.unlink(tmp_path)
        except OSError:
            pass
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        # Some .doc files are mislabeled (themeManager XML); detect and skip.
        if result.stderr and ("not a Word document" in result.stderr or
                              "is not" in result.stderr):
            return ""
        return result.stdout or ""
    except FileNotFoundError:
        # antiword not installed — degrade silently
        return ""
    except subprocess.TimeoutExpired:
        print(f"  [harvest-dl] antiword timeout", flush=True)
        return ""
    except Exception as exc:  # noqa: BLE001
        print(f"  [harvest-dl] antiword error: {type(exc).__name__}: {exc}", flush=True)
        return ""


def _docx_text(data: bytes) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return ""
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(parts)


def _xls_text(data: bytes) -> str:
    try:
        import xlrd  # type: ignore
    except ImportError:
        return ""
    try:
        book = xlrd.open_workbook(file_contents=data, formatting_info=False)
        parts: list[str] = []
        for sheet in book.sheets():
            parts.append(f"# Sheet: {sheet.name}")
            for row_idx in range(sheet.nrows):
                row = sheet.row_values(row_idx)
                cells = [str(c).strip() for c in row if str(c).strip()]
                if cells:
                    parts.append("\t".join(cells))
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        print(f"  [harvest-dl] xlrd error: {type(exc).__name__}: {exc}", flush=True)
        return ""


def _xlsx_text(data: bytes) -> str:
    try:
        from openpyxl import load_workbook  # type: ignore
    except ImportError:
        return ""
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _classify_office_member(name: str, text: str) -> str:
    """Return one of FS / NOTES / AUDITOR / OTHER based on filename + content.

    Mirrors fetch_set_financial_filings.py:pick_member.
    """
    up = name.upper()
    if any(p in up for p in ("NOTES", "หมายเหตุ")) or "notes to the financial statements" in text.lower():
        return "NOTES"
    if any(p in up for p in ("AUDITOR", "INDEPENDENT_AUDITOR", "ผู้สอบบัญชี")) or \
       "independent auditor" in text.lower() or "รายงานของผู้สอบบัญชี" in text:
        return "AUDITOR"
    if "balance sheet" in text.lower() or "งบแสดงฐานะการเงิน" in text:
        return "FS"
    return "OTHER"


# ---------------------------------------------------------------- core: process one

def process_one(client: SetNewsClient, item: dict[str, Any], state: dict[str, Any],
                dry_run: bool = False) -> str:
    """Download + extract + write one queue item. Return final status."""
    news_id = item["news_id"]
    tk = item["ticker"]
    kind = item["kind"]
    period = item["period"]
    headline = item["headline"]

    # Detail API
    try:
        detail = fetch_detail(client, news_id)
    except Exception as exc:  # noqa: BLE001
        return f"failed:detail:{type(exc).__name__}:{exc}"

    download_url = detail.get("downloadUrl")
    file_type = (detail.get("fileType") or "").upper()
    if not download_url:
        return "failed:no_download_url"

    # SHA-256 dedup BEFORE downloading (saves bandwidth; we already have it
    # in detail.content for the warm cache, but the URL is the only stable
    # identifier until we download). Best-effort: skip the network if we
    # already have a record for this news_id.
    completed = state.setdefault("completed", {})
    if news_id in completed:
        return "skipped:already_completed"

    # Download
    try:
        payload = download_bytes(client, download_url, timeout=90.0)
    except Exception as exc:  # noqa: BLE001
        return f"failed:download:{type(exc).__name__}:{exc}"

    sha = hashlib.sha256(payload).hexdigest()
    if any(v.get("sha256") == sha for v in completed.values()):
        completed[news_id] = {"sha256": sha, "kind": kind, "ts": _now_iso()}
        return "skipped:duplicate_sha"

    # Extract
    docs: list[dict[str, Any]] = []
    if kind == "MDA" or file_type == "PDF":
        text, page_count = extract_mda_pdf(payload)
        if not text:
            # Scanned PDF with no extractable text. Don't fail outright —
            # save a needs_review marker so the user can OCR later, and
            # store the source PDF metadata so the file can be re-extracted.
            print(f"  [harvest-dl] {tk} {kind} {period} -> needs_review "
                  f"(scanned PDF, no text)", flush=True)
            return f"needs_review:scanned_pdf:page_count={page_count}"
        docs.append({
            "doctype": "MDA",
            "member_filename": Path(download_url).name or f"{tk}_{period}_MDA.pdf",
            "text": text,
            "raw_bytes_len": len(payload),
            "extraction_status": "ok",
            "page_count": page_count,
            "extractor": "pypdf-v5",
            "sha256": sha,
        })
    elif kind == "FS" or file_type == "ZIP":
        members = extract_fs_zip(payload)
        for m in members:
            m["sha256"] = hashlib.sha256(payload).hexdigest()  # ZIP-level SHA for now
            m["extractor"] = "loop4-v5"
        if not members:
            return "failed:zip_no_extract"
        docs = members
    else:
        return f"failed:unknown_kind:{kind}"

    # Build the filing + entry shapes that vault_raw_writer expects.
    # vault_raw_writer._extract_docs_from_cache expects entry.raw_markdown to
    # be a dict keyed by doctype ("MDA" / "FS" / "NOTES" / "AUDITOR"), where
    # each value is a dict with {text, sha256, ...}. Build that shape from our
    # flat list of docs. Use the ZIP-level SHA for now (per-doc SHA is the
    # next iteration).
    raw_md: dict[str, dict[str, Any]] = {}
    sha = hashlib.sha256(payload).hexdigest()
    for doc in docs:
        dt_kind = doc["doctype"]
        if dt_kind in ("MDA", "FS", "NOTES", "AUDITOR"):
            raw_md[dt_kind] = {
                "text": doc["text"],
                "sha256": sha,
                "extraction_status": doc.get("extraction_status", "ok"),
                "extractor": doc.get("extractor", "loop4-v5"),
                "member_filename": doc.get("member_filename", ""),
                "raw_bytes_len": doc.get("raw_bytes_len", 0),
            }
            if doc.get("page_count") is not None:
                raw_md[dt_kind]["page_count"] = doc["page_count"]
            if doc.get("sheet_count") is not None:
                raw_md[dt_kind]["sheet_count"] = doc["sheet_count"]

    filing = {
        "tk": tk,
        "title": headline,
        "url": item.get("url", ""),
        "period": period,
        "sector": item.get("sector", "UNKNOWN"),
        "filing_id": news_id,
    }
    entry = {
        "tk": tk,
        "title": headline,
        "raw_markdown": raw_md,
        "filing_id": news_id,
    }

    if dry_run:
        return f"dry_run:{len(docs)} docs ({', '.join(d['doctype'] for d in docs)})"

    # Delegate to vault_raw_writer.
    # vault_raw_writer.DEFAULT_VAULT_ROOT points to .../1-Raw/01-Filings (the
    # filings folder itself). The project_one() function then appends
    # RAW_DIR_NAME = "01-Filings" again, producing a double-nested
    # 01-Filings/01-Filings/<DOCTYPE>/<TK>/ path. To avoid that, pass the
    # PARENT (the "1-Raw" folder) so the final path is correct.
    vault_root = vault_raw_writer.DEFAULT_VAULT_ROOT.parent
    report = vault_raw_writer.project_one(entry, filing, vault_root)
    # Treat "same sha256, skipped" as a success — the vault file already
    # exists with matching content, so nothing needs to be written.
    skipped = report.get("skipped", [])
    all_sha_skip = skipped and all("same sha256, skipped" in s for s in skipped)
    if all_sha_skip:
        completed[news_id] = {
            "sha256": sha,
            "kind": kind,
            "ts": _now_iso(),
            "files": [],  # nothing new written
        }
        return "skipped:sha_match"
    if not report.get("writes", {}).get("raw"):
        return f"failed:write:{skipped or report}"

    completed[news_id] = {
        "sha256": sha,
        "kind": kind,
        "ts": _now_iso(),
        "files": report["writes"].get("docs", []),
    }
    return "done"


def _now_iso() -> str:
    return dt.datetime.now(dt.timezone.utc).isoformat(timespec="seconds")


# ---------------------------------------------------------------- main

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description=__doc__)
    p.add_argument("--ticker", action="append", default=[], help="Limit to ticker(s)")
    p.add_argument("--limit", type=int, default=0, help="Process at most N items")
    p.add_argument("--dry-run", action="store_true", help="Don't write to vault")
    return p.parse_args()


def main() -> int:
    args = parse_args()
    queue = load_queue()
    state = load_state()

    wanted = {t.upper() for t in args.ticker}
    items = list(queue.get("items", {}).values())
    if wanted:
        items = [it for it in items if it["ticker"] in wanted]
    # Only process pending
    items = [it for it in items if it.get("status") in ("pending", "failed")]
    # Newest first
    items.sort(key=lambda it: it.get("datetime", ""), reverse=True)
    if args.limit:
        items = items[: args.limit]

    print(f"[harvest-dl] pending items: {len(items)}", flush=True)
    if not items:
        return 0

    success = failed = 0
    with SetNewsClient() as client:
        for it in items:
            print(f"[harvest-dl] {it['ticker']:6s} {it['kind']:6s} "
                  f"{it['period']:8s} #{it['news_id']} ... ", end="", flush=True)
            try:
                status = process_one(client, it, state, dry_run=args.dry_run)
            except Exception as exc:  # noqa: BLE001
                status = f"failed:exception:{type(exc).__name__}:{exc}"
            print(status, flush=True)

            if status == "done":
                it["status"] = "done"
                success += 1
            elif status.startswith("skipped"):
                it["status"] = "skipped"
                success += 1
            elif status.startswith("dry_run"):
                success += 1
            elif status.startswith("needs_review"):
                # Scanned PDF without OCR fallback available — flag for manual review.
                # Treated as success (file metadata captured, no crash) but the
                # queue item carries the caveat so the user can find it.
                it["status"] = "needs_review"
                it["needs_review_reason"] = status
                success += 1
            else:
                it["status"] = "failed"
                it["last_error"] = status
                failed += 1

    save_queue(queue)
    save_state(state)
    print(f"[harvest-dl] complete: {success} ok, {failed} failed", flush=True)
    return 0 if failed == 0 else 1


if __name__ == "__main__":
    raise SystemExit(main())
