"""Tests for the 6M26 figure extraction and the DOCX table fix.

Run: python tests/test_6m26_extraction.py

The extraction is used to build a regulator-facing panel, so most of these tests
assert what the parser *refuses* to do. A silently wrong number is far worse than
a company reported as needing review.
"""
from __future__ import annotations

import io
import sys
import unittest
from pathlib import Path

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO / "scripts"))

import extract_6m26_figures as X  # noqa: E402

FIX = REPO / "tests" / "fixtures"
CPN = (FIX / "mda_CPN_2026Q2_pl_excerpt.md").read_text(encoding="utf-8")
ITC = (FIX / "mda_ITC_2026Q2_pl_excerpt.md").read_text(encoding="utf-8")

HEADER = "Profit & Loss Statement (Baht mn) 2Q25 1Q26 2Q26 YoY (%) QoQ (%) 6M25 6M26 YoY (%)"


class TokenParsing(unittest.TestCase):
    def test_money_forms(self):
        self.assertEqual(X.parse_money("24,308"), 24308.0)
        self.assertEqual(X.parse_money("(1,806)"), -1806.0)
        self.assertEqual(X.parse_money("-450"), -450.0)
        self.assertEqual(X.parse_money("0.96"), 0.96)
        self.assertIsNone(X.parse_money("n.m."))
        self.assertIsNone(X.parse_money("-"))
        self.assertIsNone(X.parse_money("2Q25"))

    def test_percent_forms(self):
        """Issuers print negatives as (2%) and occasionally (2)%."""
        self.assertEqual(X.parse_pct("9%"), 9.0)
        self.assertEqual(X.parse_pct("(2%)"), -2.0)
        self.assertEqual(X.parse_pct("(2)%"), -2.0)
        self.assertEqual(X.parse_pct("-2.5%"), -2.5)
        self.assertIsNone(X.parse_pct("24,308"))

    def test_pdf_number_gaps_repaired(self):
        """pypdf emits '13 ,105' and '4 ,750'; both must parse as one number."""
        text = X.normalise_text("Total Revenue 13 ,105 4 ,750")
        self.assertIn("13,105", text)
        self.assertIn("4,750", text)


class ColumnGeometry(unittest.TestCase):
    def test_finds_6m_columns_past_merged_headers(self):
        """'YoY (%)' is one column printed as two tokens; mis-merging shifts everything."""
        columns = X.find_period_columns(HEADER)
        self.assertIsNotNone(columns)
        self.assertEqual((columns.prior, columns.current, columns.count), (5, 6, 8))

    def test_quarter_only_table_is_not_a_6m_table(self):
        self.assertIsNone(X.find_period_columns(
            "Statement of Financial Position (Baht mn) 2Q25 1Q26 2Q26 YoY (%) QoQ (%)"))

    def test_label_with_digits_stays_in_the_label(self):
        label, fields = X.split_row("Central Rama 2 lease income 100 200 300 5% 10% 400 500 25%")
        self.assertEqual(label, "Central Rama 2 lease income")
        self.assertEqual(len(fields), 8)


class RealFilingCPN(unittest.TestCase):
    """Values hand-checked against MDA_CPN_2026Q2_E.md in the vault."""

    def setUp(self):
        self.result = X.extract_company("CPN", CPN)

    def test_total_revenue_and_npat_reconcile(self):
        self.assertEqual(self.result.revenue.current, 26457.0)
        self.assertEqual(self.result.revenue.prior, 24308.0)
        self.assertEqual(self.result.npat.current, 9721.0)
        self.assertEqual(self.result.npat.prior, 8532.0)
        self.assertEqual(self.result.status, "verified")

    def test_rfo_derivation_matches_issuer_components(self):
        """22,957 + 1,020 + 1,247 = 25,224 — the three revenue lines in the filing."""
        self.assertEqual(self.result.rfo.current, 25224.0)
        self.assertEqual(self.result.rfo.measure, "revenue_from_operations")
        self.assertIn("26,457 - 1,233 = 25,224", self.result.rfo.derivation)

    def test_implied_margin_matches_issuer_printed_ratio(self):
        """The filing prints 6M26 net profit margin 36.7% on total revenue."""
        margin = self.result.npat.current / self.result.revenue.current * 100
        self.assertAlmostEqual(margin, 36.7, delta=0.1)

    def test_core_and_adjusted_rows_are_never_selected(self):
        for figure in (self.result.revenue, self.result.npat):
            self.assertNotIn("excl", figure.label.lower())

    def test_ratio_table_is_not_mistaken_for_a_pl_row(self):
        self.assertNotIn("margin", self.result.npat.label.lower())
        self.assertGreater(self.result.npat.current, 1000)


class RealFilingITC(unittest.TestCase):
    """A different issuer layout: 1H25/1H26 columns and '%YoY' as one token.

    Values hand-checked against MDA_ITC_2026Q2_E.md in the vault.
    """

    def setUp(self):
        self.result = X.extract_company("ITC", ITC)

    def test_reads_1h_columns_not_just_6m(self):
        self.assertEqual(self.result.rfo.prior, 8722.0)
        self.assertEqual(self.result.rfo.current, 9704.0)

    def test_sales_line_is_already_the_01_sale_basis(self):
        """'Sales and service' excludes other income, so nothing is subtracted."""
        self.assertEqual(self.result.rfo.measure, "revenue_from_operations")
        self.assertEqual(self.result.rfo.label, "Sales and service")
        self.assertEqual(self.result.rfo.derivation, "")
        self.assertEqual(self.result.other_income.current, 291.0)

    def test_normalised_sales_never_replace_reported_sales(self):
        """The filing also prints normalised 1H26 sales of 10,043."""
        self.assertNotEqual(self.result.rfo.current, 10043.0)
        self.assertNotIn("normalis", self.result.rfo.label.lower())

    def test_breakdown_table_total_is_not_taken_as_revenue(self):
        """'Total sales 8,722 10,043' in the product breakdown is normalised."""
        self.assertNotIn("total sales", self.result.rfo.label.lower())

    def test_unattributed_net_profit_is_not_promoted_to_owner_npat(self):
        self.assertFalse(self.result.npat.verified)
        self.assertTrue(self.result.npat_unattributed.verified)
        self.assertEqual(self.result.npat_unattributed.current, 1715.0)
        self.assertEqual(self.result.status, "needs_review")
        self.assertTrue(any("attributable to owners" in c for c in self.result.checks))


class RefusesToGuess(unittest.TestCase):
    def test_wrong_column_is_rejected_not_published(self):
        """If the stated YoY disagrees with the extracted pair, reject the row."""
        doc = "\n".join([
            HEADER,
            # 6M25=1,000 6M26=2,000 would be +100%, but the issuer prints 5%.
            "Total Revenue 100 200 300 5% 1% 1,000 2,000 5%",
        ])
        figure = X.extract_measure(doc, X.REVENUE_PATTERNS)
        self.assertEqual(figure.status, "unreconciled")
        self.assertIn("exceeds", figure.reason)

    def test_no_6m_table_yields_missing_not_zero(self):
        doc = "The Company reported revenue of Bt13,105mn in 2Q26, up 8% YoY."
        result = X.extract_company("XYZ", doc)
        self.assertEqual(result.status, "needs_review")
        self.assertIsNone(result.rfo.current)
        self.assertIsNone(result.npat.current)
        self.assertIn("rfo:", result.exclusion_reason)

    def test_total_revenue_without_other_income_is_not_passed_off_as_rfo(self):
        doc = "\n".join([HEADER, "Total Revenue 100 200 300 5% 1% 1,000 1,100 10%"])
        result = X.extract_company("XYZ", doc)
        self.assertTrue(result.revenue.verified)
        self.assertFalse(result.rfo.verified)
        self.assertEqual(result.status, "needs_review")

    def test_operating_profit_never_becomes_owner_npat(self):
        """An implausible margin means a mis-picked row; reject rather than report."""
        doc = "\n".join([
            HEADER,
            "Total Revenue 10 20 30 5% 1% 100 110 10%",
            "Other Income 1 1 1 0% 0% 10 11 10%",
            "Profit to Parent Company 10 20 30 5% 1% 900 990 10%",
        ])
        result = X.extract_company("XYZ", doc)
        self.assertEqual(result.npat.status, "unreconciled")
        self.assertIn("out of range", result.npat.reason)
        self.assertEqual(result.status, "needs_review")


class LossMakingIssuer(unittest.TestCase):
    """Loss-makers print no usable YoY but belong in the panel via npat_state."""

    DOC = "\n".join([
        HEADER,
        "Total Revenue 500 600 700 5% 1% 1,000 1,200 20%",
        "Other Income 10 10 10 0% 0% 20 24 20%",
        "Profit to Parent Company (100) (120) (90) 5% 1% (800) (450) n.m.",
    ])

    def test_negative_base_accepted_via_proven_column_geometry(self):
        result = X.extract_company("XYZ", self.DOC)
        self.assertEqual(result.npat.current, -450.0)
        self.assertEqual(result.npat.prior, -800.0)
        self.assertTrue(result.npat.verified)
        self.assertIn("column mapping proven", result.npat.reason)

    def test_geometry_never_rescues_a_contradicted_row(self):
        """A row whose own YoY disagrees is a wrong pick, not a missing YoY."""
        contradicted = X.Figure(status="unreconciled", header=HEADER,
                                reason="stated YoY +5.0% vs computed +100.0% exceeds 1.0pp",
                                prior=1.0, current=2.0)
        proven = X.Figure(status="verified", header=HEADER, label="Total Revenue")
        self.assertFalse(X.reconcile_by_geometry(contradicted, proven).verified)

    def test_geometry_requires_the_same_table(self):
        other = X.Figure(status="unreconciled", header="A DIFFERENT TABLE 6M25 6M26",
                         reason="no issuer-stated YoY beside the 6M26 column",
                         prior=1.0, current=2.0)
        proven = X.Figure(status="verified", header=HEADER, label="Total Revenue")
        self.assertFalse(X.reconcile_by_geometry(other, proven).verified)


class DocxTableExtraction(unittest.TestCase):
    """harvest_download._docx_text dropped every table before 2026-08-17."""

    def setUp(self):
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("python-docx not installed")
        import harvest_download  # noqa: E402
        self.harvest = harvest_download

    def _build_docx(self) -> bytes:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Revenue and results for the six-month period ended 30 June:")
        table = doc.add_table(rows=2, cols=3)
        for column, value in enumerate(["Item", "6M25", "6M26"]):
            table.cell(0, column).text = value
        for column, value in enumerate(["Total Revenue", "24,308", "26,457"]):
            table.cell(1, column).text = value
        doc.add_paragraph("These notes form an integral part of the statements.")
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def test_table_numbers_survive_extraction(self):
        text = self.harvest._docx_text(self._build_docx())
        self.assertIn("24,308", text)
        self.assertIn("26,457", text)

    def test_table_stays_in_document_order(self):
        text = self.harvest._docx_text(self._build_docx())
        intro = text.index("Revenue and results")
        row = text.index("24,308")
        outro = text.index("These notes form")
        self.assertLess(intro, row, "table must follow the heading that introduces it")
        self.assertLess(row, outro, "table must precede the paragraph after it")

    def test_row_is_tab_separated_like_the_xls_path(self):
        text = self.harvest._docx_text(self._build_docx())
        self.assertIn("Total Revenue\t24,308\t26,457", text)

    def test_merged_cells_are_not_duplicated(self):
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        table.cell(0, 0).text = "Merged header"
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 2).text = "6M26"
        buffer = io.BytesIO()
        doc.save(buffer)
        text = self.harvest._docx_text(buffer.getvalue())
        self.assertEqual(text.count("Merged header"), 1)

    def test_extracted_table_feeds_the_parser(self):
        """End to end: the fix makes a DOCX filing extractable."""
        from docx import Document
        doc = Document()
        table = doc.add_table(rows=3, cols=9)
        rows = [
            ["Profit & Loss (Baht mn)", "2Q25", "1Q26", "2Q26", "YoY (%)", "QoQ (%)", "6M25", "6M26", "YoY (%)"],
            ["Total Revenue", "10", "20", "30", "5%", "1%", "1,000", "1,100", "10%"],
            ["Other Income", "1", "1", "1", "0%", "0%", "100", "110", "10%"],
        ]
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                table.cell(r, c).text = value
        buffer = io.BytesIO()
        doc.save(buffer)
        text = self.harvest._docx_text(buffer.getvalue()).replace("\t", " ")
        result = X.extract_company("XYZ", text)
        self.assertEqual(result.rfo.current, 990.0)
        self.assertTrue(result.rfo.verified)


if __name__ == "__main__":
    unittest.main(verbosity=2)
