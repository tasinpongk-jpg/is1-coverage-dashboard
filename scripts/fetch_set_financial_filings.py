#!/usr/bin/env python3
"""Fetch SET annual financial-statement filing ZIPs and write vault markdown.

The 2025 annual SET filing downloads for some issuers are Office bundles rather
than PDFs. This helper downloads those ZIPs, extracts DOCX text, and writes the
FS-NOTES/AUDITOR markdown files consumed by build_vault_ticker_notes.py.
"""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import os
import re
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import httpx
from docx import Document
from openpyxl import load_workbook

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from surveillance.client import HEADERS  # noqa: E402

REPO = Path(__file__).resolve().parent.parent
DEFAULT_VAULT = Path.home() / "Library" / "CloudStorage" / "OneDrive2-TheStockExchangeofThailand" / "Claude-Vault"
LISTED_ROOT = Path("Work-SET") / "Listed Company"
FS_NOTES_SUBPATH = Path("1-Raw") / "01-Filings" / "FS-NOTES"
AUDITOR_SUBPATH = Path("1-Raw") / "01-Filings" / "AUDITOR"


@dataclass(frozen=True)
class FilingJob:
    ticker: str
    news_id: str
    zip_url: str
    period: str
    fy_end_date: str
    needs: tuple[str, ...]


JOBS = [
    FilingJob("FPT", "17624721958960", "https://weblink.set.or.th/dat/news/202511/0675FIN071120251059422480E.zip", "2025FY", "2025-09-30", ("NOTES", "AUDITOR")),
    FilingJob("KSL", "17661013124860", "https://weblink.set.or.th/dat/news/202512/0828FIN191220252042550687E.zip", "2025FY", "2025-10-31", ("NOTES", "AUDITOR")),
    FilingJob("KTIS", "17642869843950", "https://weblink.set.or.th/dat/news/202511/1149FIN281120252045060593E.zip", "2025FY", "2025-09-30", ("NOTES", "AUDITOR")),
    FilingJob("UV", "17639414126340", "https://weblink.set.or.th/dat/news/202511/0136FIN241120251726190547E.zip", "2025FY", "2025-09-30", ("NOTES", "AUDITOR")),
    FilingJob("PREB", "17714553188280", "https://weblink.set.or.th/dat/news/202602/0871FIN190220262141330673E.zip", "2025FY", "2025-12-31", ("NOTES",)),
    FilingJob("TPOLY", "17721442966690", "https://weblink.set.or.th/dat/news/202602/0988FIN270220260845286310E.zip", "2025FY", "2025-12-31", ("NOTES",)),
]


def normalize_url(url: str) -> str:
    if url.startswith("http"):
        return url
    return "https://" + url.lstrip("/")


def clean_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    return text.strip()


def docx_text(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append("\t".join(cells))
    return clean_text("\n\n".join(parts))


def legacy_doc_text(data: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".doc") as tmp:
        tmp.write(data)
        tmp.flush()
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-stdout", tmp.name],
            check=True,
            capture_output=True,
            text=True,
        )
    return clean_text(result.stdout)


def xlsx_text(data: bytes) -> str:
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append("\t".join(cells))
    return clean_text("\n".join(parts))


def iter_zip_members(data: bytes, prefix: str = "") -> Iterable[tuple[str, bytes]]:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            name = prefix + info.filename
            payload = zf.read(info)
            suffix = Path(name).suffix.lower()
            if suffix in {".doc", ".docx", ".xls", ".xlsx"}:
                yield name, payload
            elif zipfile.is_zipfile(io.BytesIO(payload)):
                yield from iter_zip_members(payload, prefix=f"{name}/")
            else:
                yield name, payload


def extract_texts(zip_bytes: bytes) -> dict[str, tuple[str, bytes, str]]:
    out: dict[str, tuple[str, bytes, str]] = {}
    for name, payload in iter_zip_members(zip_bytes):
        suffix = Path(name).suffix.lower()
        try:
            if suffix in {".doc", ".docx"}:
                try:
                    text = docx_text(payload)
                except Exception:
                    if suffix == ".doc":
                        text = legacy_doc_text(payload)
                    else:
                        raise
            elif suffix in {".xls", ".xlsx"}:
                text = xlsx_text(payload)
            elif suffix in {".txt", ".csv", ".xml"}:
                text = payload.decode("utf-8", errors="replace")
            else:
                continue
        except Exception as exc:
            print(f"  warn: failed to parse {name}: {exc}")
            continue
        out[name.upper()] = (text, payload, name)
    return out


def pick_member(texts: dict[str, tuple[str, bytes, str]], kind: str) -> tuple[str, bytes, str]:
    patterns = {
        "NOTES": ("NOTES", "หมายเหตุ"),
        "AUDITOR": ("AUDITOR", "INDEPENDENT_AUDITOR", "ผู้สอบบัญชี"),
    }[kind]
    for key, value in texts.items():
        if any(p in key for p in patterns):
            return value
    haystacks = {
        "NOTES": ("notes to the financial statements", "หมายเหตุประกอบงบการเงิน"),
        "AUDITOR": ("independent auditor", "รายงานของผู้สอบบัญชี"),
    }[kind]
    for value in texts.values():
        text = value[0].lower()
        if any(p.lower() in text for p in haystacks):
            return value
    raise KeyError(f"could not find {kind} member")


def trim_to_section(text: str, kind: str) -> str:
    starts = {
        "NOTES": (
            r"(?im)^notes?\s+to\s+the\s+(?:consolidated\s+and\s+separate\s+)?financial\s+statements",
            r"(?im)^หมายเหตุประกอบงบการเงิน",
            r"(?im)^1\s+general information",
        ),
        "AUDITOR": (
            r"(?im)^independent auditor'?s report",
            r"(?im)^รายงานของผู้สอบบัญชี",
        ),
    }[kind]
    for pat in starts:
        match = re.search(pat, text)
        if match:
            return text[match.start() :].strip()
    return text.strip()


def markdown(job: FilingJob, kind: str, text: str, source_name: str, source_bytes: bytes) -> str:
    filing_type = "FS_NOTES" if kind == "NOTES" else "AUDITOR"
    title = "FS_NOTES" if kind == "NOTES" else "AUDITOR"
    tag = "fs_notes" if kind == "NOTES" else "auditor"
    sha = hashlib.sha256(source_bytes).hexdigest()
    body = trim_to_section(clean_text(text), kind)
    frontmatter = {
        "ticker": job.ticker,
        "filing_type": filing_type,
        "period": job.period,
        "period_kind": "annual",
        "language": "E",
        "fy_end_date": job.fy_end_date,
        "news_id": job.news_id,
        "source_sha256": sha,
        "source_url": job.zip_url,
        "source_file": source_name,
        "tags": f"[filing, {tag}, ticker/{job.ticker}]",
        "source_type": "set_zip_office",
    }
    lines = ["---"]
    for key, value in frontmatter.items():
        if key in {"source_url", "source_file"}:
            lines.append(f'{key}: "{value}"')
        else:
            lines.append(f"{key}: {value}")
    lines.extend(
        [
            "---",
            f"# {title} — {job.ticker} {job.period} (E)",
            f"_Source: `{source_name}` from SET news `{job.news_id}` · SHA-256 `{sha[:12]}...`_",
            "",
            body,
            "",
        ]
    )
    return "\n".join(lines)


def output_path(vault: Path, job: FilingJob, kind: str) -> Path:
    if kind == "NOTES":
        return vault / LISTED_ROOT / FS_NOTES_SUBPATH / job.ticker / f"NOTES_{job.ticker}_{job.period}_E.md"
    return vault / LISTED_ROOT / AUDITOR_SUBPATH / job.ticker / f"AUDITOR_{job.ticker}_{job.period}_E.md"


def download(client: httpx.Client, job: FilingJob) -> bytes:
    detail_url = f"https://www.set.or.th/api/set/news/{job.news_id}/detail"
    try:
        detail = client.get(detail_url, timeout=30)
        if detail.status_code == 200:
            download_url = detail.json().get("downloadUrl")
            if download_url:
                job_url = normalize_url(str(download_url))
                if job_url != job.zip_url:
                    print(f"  detail downloadUrl: {job_url}")
    except Exception as exc:
        print(f"  warn: detail API failed for {job.ticker}: {exc}")
    response = client.get(job.zip_url, timeout=60)
    response.raise_for_status()
    return response.content


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--vault-root", default=os.environ.get("VAULT_ROOT", str(DEFAULT_VAULT)))
    parser.add_argument("--ticker", action="append", default=[], help="Only process selected ticker(s).")
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args(argv)

    vault = Path(args.vault_root).expanduser()
    wanted = {t.upper() for t in args.ticker}
    jobs = [j for j in JOBS if not wanted or j.ticker in wanted]

    headers = dict(HEADERS)
    headers["Referer"] = "https://www.set.or.th/"
    with httpx.Client(headers=headers, follow_redirects=True) as client:
        client.get("https://www.set.or.th", timeout=30).raise_for_status()
        for job in jobs:
            print(f"{job.ticker}: downloading {job.zip_url}")
            zip_bytes = download(client, job)
            print(f"  zip bytes={len(zip_bytes)} sha={hashlib.sha256(zip_bytes).hexdigest()[:12]}")
            texts = extract_texts(zip_bytes)
            print("  members:", ", ".join(v[2] for v in texts.values()))
            for kind in job.needs:
                text, payload, source_name = pick_member(texts, kind)
                out_path = output_path(vault, job, kind)
                content = markdown(job, kind, text, source_name, payload)
                print(f"  write {out_path} ({len(content):,} chars)")
                if not args.dry_run:
                    out_path.parent.mkdir(parents=True, exist_ok=True)
                    out_path.write_text(content, encoding="utf-8")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
